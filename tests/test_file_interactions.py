from __future__ import annotations

from pathlib import Path

from fastapi.testclient import TestClient

from agentflow.agents.answer.agent import AnswerAgent
from agentflow.app.main import app


def test_create_and_read_generated_file_preserves_nested_path(tmp_path: Path):
    client = TestClient(app)
    workspace = tmp_path / "workspace"
    workspace.mkdir()

    set_response = client.post("/workspace/set", json={"path": str(workspace)})
    assert set_response.status_code == 200

    create_response = client.post(
        "/files/create",
        json={
            "filename": "snake_game/python_snake.py",
            "content": "print('snake')\n",
            "workspace_path": str(workspace),
        },
    )
    assert create_response.status_code == 200
    create_body = create_response.json()
    assert create_body["filename"] == "snake_game/python_snake.py"
    assert (workspace / "snake_game" / "python_snake.py").exists()

    read_response = client.post(
        "/files/read",
        json={"path": create_body["path"], "workspace_path": str(workspace)},
    )
    assert read_response.status_code == 200
    read_body = read_response.json()
    assert read_body["content"].replace("\r\n", "\n") == "print('snake')\n"
    assert read_body["truncated"] is False


def test_read_generated_docx_extracts_text_preview(tmp_path: Path):
    from docx import Document

    client = TestClient(app)
    workspace = tmp_path / "workspace"
    workspace.mkdir()

    set_response = client.post("/workspace/set", json={"path": str(workspace)})
    assert set_response.status_code == 200

    docx_path = workspace / "report.docx"
    doc = Document()
    doc.add_heading("OmniForge 报告", level=1)
    doc.add_paragraph("这是 Word 文档预览内容。")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "能力"
    table.rows[0].cells[1].text = "知识库"
    doc.save(str(docx_path))

    read_response = client.post(
        "/files/read",
        json={"path": str(docx_path), "workspace_path": str(workspace)},
    )

    assert read_response.status_code == 200
    body = read_response.json()
    assert "OmniForge 报告" in body["content"]
    assert "这是 Word 文档预览内容。" in body["content"]
    assert "能力 | 知识库" in body["content"]
    assert "PK" not in body["content"][:10]


def test_completion_summary_does_not_report_directories_as_files():
    state = {
        "goal_analysis": {
            "goal": "创建两个文件",
            "goal_type": "project",
        },
        "plan": {},
        "task_queue": [
            {
                "status": "done",
                "tool": "filesystem",
                "goal": "mkdir",
                "input": {"action": "mkdir", "path": "snake_game"},
            }
        ],
    }

    result = AnswerAgent().run(state)

    assert "snake_game" not in result["answer"]
