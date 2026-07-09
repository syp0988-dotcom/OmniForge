"""DocxTool — Word document creation, reading, and editing.

Uses python-docx for high-level operations (create, read, edit) and
the Claude Code docx skill scripts for low-level XML manipulation
(unpack, pack, validate, tracked-changes).

Requirements:
    python-docx  (installed)
    defusedxml   (for skill scripts — pip install defusedxml)
"""

from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path
from typing import Any

from agentflow.tools.base import BaseTool
from agentflow.tools.result import ToolResult
from agentflow.utils.logging import build_logger

logger = build_logger("docx_tool")

# -- Resolve skill scripts path ------------------------------------------------

_SKILL_SCRIPTS = Path(
    os.path.expandvars(r"%USERPROFILE%\.claude\skills\docx\scripts")
)
_SCRIPTS_PATH_ADDED = False


def _ensure_skill_path() -> None:
    """Add skill scripts to sys.path once."""
    global _SCRIPTS_PATH_ADDED
    if not _SCRIPTS_PATH_ADDED:
        sp = str(_SKILL_SCRIPTS)
        if sp not in sys.path:
            sys.path.insert(0, sp)
        office_sp = str(_SKILL_SCRIPTS / "office")
        if office_sp not in sys.path:
            sys.path.insert(0, office_sp)
        _SCRIPTS_PATH_ADDED = True


# ---------------------------------------------------------------------------
# DocxTool
# ---------------------------------------------------------------------------


class DocxTool(BaseTool):
    """Create, read, and edit Word (.docx) documents.

    High-level operations use python-docx; low-level XML manipulation
    delegates to the Claude Code docx skill scripts.
    """

    name = "docx"
    description = "Word 文档操作 — 创建、读取、编辑、添加批注、验证 .docx 文件"

    def __init__(self, workspace: str = ".") -> None:
        self.workspace = Path(workspace).resolve()

    def _resolve(self, path: str) -> Path:
        p = Path(path)
        if p.is_absolute():
            return p
        return self.workspace / p

    # ------------------------------------------------------------------
    # Introspection
    # ------------------------------------------------------------------

    def actions(self) -> dict[str, dict]:
        return {
            "create": {
                "description": (
                    "创建新的 Word 文档（.docx）。使用 Markdown 风格的 content 描述文档结构，"
                    "支持 # 标题、普通段落、表格（用 | 分隔列，每行一个）、无序列表（- 开头）。"
                    "示例：\n"
                    '"# 报告标题\\n\\n## 第一节\\n\\n这是段落内容。\\n\\n| 列1 | 列2 |\\n| 数据1 | 数据2 |"'
                ),
                "parameters": {
                    "path": {"type": "string", "description": "输出 .docx 文件路径"},
                    "content": {"type": "string", "description": "Markdown 风格的文档内容描述"},
                },
                "required": ["path", "content"],
            },
            "read": {
                "description": "读取 .docx 文件的文本内容，返回所有段落文本",
                "parameters": {
                    "path": {"type": "string", "description": ".docx 文件路径"},
                },
                "required": ["path"],
            },
            "read_tables": {
                "description": "读取 .docx 文件中的所有表格内容",
                "parameters": {
                    "path": {"type": "string", "description": ".docx 文件路径"},
                },
                "required": ["path"],
            },
            "replace_text": {
                "description": "在 .docx 文件中查找并替换文本（支持所有段落和表格）",
                "parameters": {
                    "path": {"type": "string", "description": ".docx 文件路径"},
                    "old_text": {"type": "string", "description": "要查找的文本"},
                    "new_text": {"type": "string", "description": "替换后的文本"},
                },
                "required": ["path", "old_text", "new_text"],
            },
            "add_comment": {
                "description": (
                    "向 .docx 文件的指定段落添加批注。"
                    "paragraph_index 从 0 开始计数"
                ),
                "parameters": {
                    "path": {"type": "string", "description": ".docx 文件路径"},
                    "paragraph_index": {
                        "type": "integer",
                        "description": "要添加批注的段落索引（从 0 开始）",
                    },
                    "text": {"type": "string", "description": "批注文本内容"},
                    "author": {
                        "type": "string",
                        "description": "批注作者名称（默认：Reviewer）",
                    },
                },
                "required": ["path", "paragraph_index", "text"],
            },
            "validate": {
                "description": "验证 .docx 文件结构是否合法（使用 XSD 模式校验）",
                "parameters": {
                    "path": {"type": "string", "description": ".docx 文件路径"},
                },
                "required": ["path"],
            },
            "convert_to_pdf": {
                "description": (
                    "将 .docx 转换为 PDF（需要安装 LibreOffice）。"
                    "转换后的 PDF 保存在同一目录下"
                ),
                "parameters": {
                    "path": {"type": "string", "description": ".docx 文件路径"},
                },
                "required": ["path"],
            },
        }

    # ------------------------------------------------------------------
    # Execute
    # ------------------------------------------------------------------

    def execute(self, action: str = "", **kwargs: Any) -> ToolResult:
        if action not in self.actions():
            return ToolResult.fail(
                self.name, action or "execute",
                f"Unknown action '{action}'. Available: {list(self.actions())}",
            )
        handler = getattr(self, f"_cmd_{action}", None)
        if handler is None:
            return ToolResult.fail(self.name, action, f"No handler for '{action}'")
        try:
            return handler(**kwargs)
        except Exception as exc:
            logger.exception("DocxTool.%s failed", action)
            return ToolResult.fail(self.name, action, str(exc))

    # ==================================================================
    # Action handlers
    # ==================================================================

    # -- create ---------------------------------------------------------

    def _cmd_create(self, path: str = "", content: str = "", **kwargs: Any) -> ToolResult:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        filepath = self._resolve(path)

        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)

        # Parse markdown-style content
        lines = content.split("\n")
        i = 0
        paragraph_count = 0
        table_count = 0

        while i < len(lines):
            line = lines[i]

            # Table: lines starting with | (collect consecutive | lines)
            if line.strip().startswith("|") and line.strip().endswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1

                if len(table_lines) >= 2:
                    # First row = header, skip separator row if present
                    header_cells = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                    data_start = 1
                    # Skip separator like |---|---|
                    if data_start < len(table_lines):
                        sep_cells = [c.strip() for c in table_lines[1].split("|")[1:-1]]
                        if all(set(c) <= {"-", ":", " "} for c in sep_cells):
                            data_start = 2

                    num_cols = len(header_cells)
                    table = doc.add_table(rows=1 + len(table_lines) - data_start, cols=num_cols)
                    table.style = "Table Grid"

                    # Header
                    for ci, text in enumerate(header_cells):
                        cell = table.rows[0].cells[ci]
                        cell.text = text
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.bold = True

                    # Data rows
                    for ri, row_line in enumerate(table_lines[data_start:], start=1):
                        cells = [c.strip() for c in row_line.split("|")[1:-1]]
                        for ci, text in enumerate(cells[:num_cols]):
                            table.rows[ri].cells[ci].text = text

                    table_count += 1
                continue

            # Heading
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
                paragraph_count += 1
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
                paragraph_count += 1
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
                paragraph_count += 1
            # Unordered list
            elif line.strip().startswith("- "):
                doc.add_paragraph(line.strip()[2:], style="List Bullet")
                paragraph_count += 1
            # Ordered list
            elif line.strip() and line.strip()[0].isdigit() and ". " in line.strip()[:4]:
                text = line.strip().split(". ", 1)[1] if ". " in line.strip() else line.strip()
                doc.add_paragraph(text, style="List Number")
                paragraph_count += 1
            # Empty line → skip
            elif not line.strip():
                pass
            # Regular paragraph
            else:
                doc.add_paragraph(line)
                paragraph_count += 1

            i += 1

        filepath.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(filepath))

        return ToolResult.ok(
            self.name, "create",
            result={
                "path": str(filepath),
                "paragraphs": paragraph_count,
                "tables": table_count,
            },
            message=f"Created {filepath.name}: {paragraph_count} paragraphs, {table_count} tables",
        )

    # -- read -----------------------------------------------------------

    def _cmd_read(self, path: str = "", **kwargs: Any) -> ToolResult:
        from docx import Document

        filepath = self._resolve(path)
        if not filepath.exists():
            return ToolResult.fail(self.name, "read", f"File not found: {filepath}")

        doc = Document(str(filepath))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = "\n".join(paragraphs)

        return ToolResult.ok(
            self.name, "read",
            result={
                "path": str(filepath),
                "paragraph_count": len(paragraphs),
                "text": full_text,
                "paragraphs": paragraphs,
            },
            message=f"Read {len(paragraphs)} paragraphs from {filepath.name}",
        )

    # -- read_tables ----------------------------------------------------

    def _cmd_read_tables(self, path: str = "", **kwargs: Any) -> ToolResult:
        from docx import Document

        filepath = self._resolve(path)
        if not filepath.exists():
            return ToolResult.fail(self.name, "read_tables", f"File not found: {filepath}")

        doc = Document(str(filepath))
        tables_data = []
        for ti, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables_data.append({"table_index": ti, "rows": rows, "row_count": len(rows)})

        return ToolResult.ok(
            self.name, "read_tables",
            result={"path": str(filepath), "tables": tables_data, "table_count": len(tables_data)},
            message=f"Read {len(tables_data)} tables from {filepath.name}",
        )

    # -- replace_text ---------------------------------------------------

    def _cmd_replace_text(self, path: str = "", old_text: str = "", new_text: str = "", **kwargs: Any) -> ToolResult:
        from docx import Document

        filepath = self._resolve(path)
        if not filepath.exists():
            return ToolResult.fail(self.name, "replace_text", f"File not found: {filepath}")

        doc = Document(str(filepath))
        replaced = 0

        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                # Inline runs need per-run replacement
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        replaced += 1

        # Also search tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_text in paragraph.text:
                            for run in paragraph.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
                                    replaced += 1

        doc.save(str(filepath))
        return ToolResult.ok(
            self.name, "replace_text",
            result={"path": str(filepath), "occurrences_replaced": replaced},
            message=f"Replaced '{old_text}' → '{new_text}' in {replaced} occurrence(s)",
        )

    # -- add_comment ----------------------------------------------------

    def _cmd_add_comment(self, path: str = "", paragraph_index: int = 0, text: str = "", author: str = "Reviewer", **kwargs: Any) -> ToolResult:
        from docx import Document

        filepath = self._resolve(path)
        if not filepath.exists():
            return ToolResult.fail(self.name, "add_comment", f"File not found: {filepath}")

        doc = Document(str(filepath))
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return ToolResult.fail(
                self.name, "add_comment",
                f"Paragraph index {paragraph_index} out of range (0-{len(doc.paragraphs) - 1})",
            )

        para = doc.paragraphs[paragraph_index]
        # python-docx 1.2.0: add_comment on Document
        try:
            doc.add_comment(text, author=author, initials=author[:2].upper(), range=para)
            doc.save(str(filepath))
            return ToolResult.ok(
                self.name, "add_comment",
                result={
                    "path": str(filepath),
                    "paragraph_index": paragraph_index,
                    "author": author,
                    "paragraph_text": para.text[:100],
                },
                message=f"Comment added to paragraph {paragraph_index} by '{author}'",
            )
        except TypeError:
            # Fallback: python-docx API may differ slightly
            return ToolResult.fail(
                self.name, "add_comment",
                "Comment API not fully supported in this python-docx version. "
                "Try using the docx skill scripts directly.",
            )

    # -- validate -------------------------------------------------------

    def _cmd_validate(self, path: str = "", **kwargs: Any) -> ToolResult:
        _ensure_skill_path()
        from office.validate import main as validate_main

        filepath = self._resolve(path)
        if not filepath.exists():
            return ToolResult.fail(self.name, "validate", f"File not found: {filepath}")

        try:
            validate_main([str(filepath), "--auto-repair", "--author", "DocxTool"])
            return ToolResult.ok(
                self.name, "validate",
                result={"path": str(filepath), "valid": True},
                message=f"Validation passed: {filepath.name}",
            )
        except SystemExit as e:
            if e.code == 0 or e.code is None:
                return ToolResult.ok(
                    self.name, "validate",
                    result={"path": str(filepath), "valid": True},
                    message=f"Validation passed: {filepath.name}",
                )
            return ToolResult.fail(
                self.name, "validate",
                f"Validation failed with code {e.code}",
                result={"path": str(filepath), "valid": False, "exit_code": e.code},
            )

    # -- convert_to_pdf -------------------------------------------------

    def _cmd_convert_to_pdf(self, path: str = "", **kwargs: Any) -> ToolResult:
        import subprocess as sp

        filepath = self._resolve(path)
        if not filepath.exists():
            return ToolResult.fail(self.name, "convert_to_pdf", f"File not found: {filepath}")

        output_dir = str(filepath.parent)
        try:
            result = sp.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, str(filepath)],
                capture_output=True, text=True, timeout=60,
            )
            pdf_path = filepath.with_suffix(".pdf")
            if pdf_path.exists():
                return ToolResult.ok(
                    self.name, "convert_to_pdf",
                    result={"path": str(pdf_path), "docx_path": str(filepath)},
                    message=f"Converted to PDF: {pdf_path.name}",
                )
            return ToolResult.fail(
                self.name, "convert_to_pdf",
                f"PDF not created. soffice stdout: {result.stdout[:200]}, stderr: {result.stderr[:200]}",
            )
        except FileNotFoundError:
            return ToolResult.fail(
                self.name, "convert_to_pdf",
                "LibreOffice not found. Install LibreOffice to enable docx→PDF conversion.",
            )
        except sp.TimeoutExpired:
            return ToolResult.fail(self.name, "convert_to_pdf", "Conversion timed out (60s)")
