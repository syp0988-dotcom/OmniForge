"""API routes for chat and knowledge base management."""

from __future__ import annotations

import asyncio
import io
import json
import shutil
from datetime import datetime
from pathlib import Path
from tempfile import NamedTemporaryFile

from pydantic import BaseModel

from fastapi import APIRouter, File, HTTPException, Request, UploadFile
from fastapi.responses import JSONResponse, StreamingResponse

from agentflow.agents.registry import get_all as get_all_agents
from agentflow.database.sqlite import SQLiteStore
from agentflow.graph.workflow import build_workflow, run_workflow
from agentflow.knowledge.store import KnowledgeStore
from agentflow.models.chat import ChatRequest, ChatResponse, FileProposal
from agentflow.models.model_config import LLMModelCreate, LLMModelUpdate
from agentflow.services.file_proposer import propose_files
from agentflow.utils.logging import build_logger

router = APIRouter()
logger = build_logger("api")

# -- Lazy-initialised store accessors (decoupled for testability) -------------
# Replace module-level globals with lazy accessors so tests can swap
# implementations by calling set_store(mock_store) before routes are hit.

_store: SQLiteStore | None = None
_knowledge_store: KnowledgeStore | None = None
_workspace_root: Path | None = None


def get_store() -> SQLiteStore:
    global _store
    if _store is None:
        _store = SQLiteStore()
    return _store


def get_knowledge_store() -> KnowledgeStore:
    global _knowledge_store
    if _knowledge_store is None:
        _knowledge_store = KnowledgeStore(db=get_store())
    return _knowledge_store


def set_store(store: SQLiteStore) -> None:
    """Override the global store (for testing / DI)."""
    global _store
    _store = store


def set_knowledge_store(ks: KnowledgeStore) -> None:
    """Override the global knowledge store (for testing / DI)."""
    global _knowledge_store
    _knowledge_store = ks

# Directory for uploaded document files
UPLOAD_DIR = Path(__file__).resolve().parents[2] / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

OUTPUT_DIR = Path(__file__).resolve().parents[2] / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)


def _is_relative_to(path: Path, parent: Path) -> bool:
    """Return True when ``path`` is inside ``parent`` after resolution."""
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except ValueError:
        return False


def _current_workspace_root() -> Path:
    """Return the active workspace root, defaulting to the app project root."""
    return (_workspace_root or Path(__file__).resolve().parents[2]).resolve()


def _resolve_workspace_child(path: str | None, *, default: Path | None = None) -> Path:
    """Resolve a path and require it to stay inside the active workspace root."""
    root = _current_workspace_root()
    target = (default or root) if not path else Path(path).resolve()
    if not _is_relative_to(target, root):
        raise HTTPException(status_code=403, detail="Path is outside the active workspace")
    return target


# -- Agent introspection ---------------------------------------------------


@router.get("/agents")
def list_agents() -> list[dict[str, object]]:
    """List all registered agents with metadata (name, key, status, capabilities)."""
    return get_all_agents()


# -- Chat -------------------------------------------------------------------


@router.post("/chat")
async def chat(request: ChatRequest):
    """Handle chat requests — delegates to the streaming generator internally.

    This endpoint is non-blocking (async) and streams the workflow execution
    in the background, returning only the final result after completion.
    Previously this was a blocking synchronous call — now it uses the same
    async infrastructure as ``/chat/stream`` for consistency.
    """
    from agentflow.conversation.session_state import SessionState
    from agentflow.graph.context import WorkflowContext

    # -- Session setup --
    session_id = request.session_id
    if session_id is None:
        sess = get_store().create_session()
        session_id = sess["id"]
    else:
        existing = get_store().get_session(session_id)
        if existing is None:
            raise HTTPException(status_code=404, detail="Session not found")

    history_dicts = [
        {"role": m.role, "content": m.content} for m in request.history
    ]

    # Load session_state from DB
    saved_state_str = get_store().get_session_state(session_id)
    session_state_dict = json.loads(saved_state_str) if saved_state_str else None

    workflow = build_workflow()

    initial_state: dict = {
        "question": request.message,
        "workflow": [],
        "history": history_dicts,
        "source_mode": request.source_mode,
    }
    if history_dicts:
        initial_state["memory"] = {"history": list(history_dicts)}
    if session_state_dict:
        initial_state["session_state"] = SessionState.from_dict(session_state_dict)

    # Use astream to capture all node outputs (same approach as test_debug.py).
    # Each node in our graph returns the full state dict, so we accumulate
    # the latest output from each node as the graph progresses.
    try:
        final_state: dict | None = None
        async for event in workflow.astream(initial_state):
            for node_name, state_update in event.items():
                # Each node returns the full state dict — capture the latest.
                final_state = dict(state_update)
                logger.debug("Node '%s' emitted keys: %s", node_name, list(state_update.keys())[:10])

        # After all nodes complete, final_state should hold the state after "memory" (the end node).
        answer = (final_state or {}).get("answer", "")
        error = (final_state or {}).get("error", "")

        logger.info("Chat final_state keys: %s", list((final_state or {}).keys())[:20])
        if not answer:
            logger.warning("Chat: answer is empty. error=%s keys=%s",
                           error, list((final_state or {}).keys())[:20])

        # Persist messages
        get_store().add_message("user", request.message, session_id=session_id)
        if answer:
            get_store().add_message("assistant", answer, session_id=session_id)

        # Persist session_state
        if final_state:
            ctx = WorkflowContext(final_state)
            result_dict = ctx.to_dict()
            new_state = result_dict.get("session_state")
            if new_state and isinstance(new_state, dict):
                get_store().update_session_state(session_id, json.dumps(new_state, ensure_ascii=False))

        # Auto-title
        sess = get_store().get_session(session_id)
        if sess and sess["title"] == "新对话":
            title = request.message[:50]
            if len(request.message) > 50:
                title += "…"
            get_store().update_session_title(session_id, title)

        debug_data = {
            "goal": (final_state or {}).get("goal_analysis", {}),
            "category": (final_state or {}).get("category"),
            "workflow": (final_state or {}).get("workflow"),
            "search_results": (final_state or {}).get("search_results", []),
        }
        return ChatResponse(
            reply=answer,
            metadata={"status": "ok", "session_id": session_id},
            debug=debug_data,
            proposed_files=propose_files(answer),
        )
    except Exception as exc:
        logger.exception("Chat error")
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@router.post("/chat/stream")
async def chat_stream(body: ChatRequest, raw_request: Request):
    """SSE streaming endpoint — yields events as workflow nodes complete.

    Now supports client-disconnect detection: when the frontend AbortController
    fires, the backend detects ``request.is_disconnected()`` between workflow
    nodes and stops early, skipping persistence and sending a ``cancelled``
    SSE event.

    Events::

        event: thinking\\n data: {{"phase": "分析问题", "category": "search"}}
        event: searching\\n data: {{"phase": "搜索网络信息"}}
        event: generating\\n data: {{"phase": "生成回答"}}
        event: cancelled\\n data: {{"reason": "用户中断了对话"}}
        event: done\\n data: {{"answer": "..."}}
    """
    from agentflow.conversation.session_state import SessionState
    from agentflow.graph.context import WorkflowContext

    async def _event_generator():
        # -- Helper: check client disconnect --
        async def _is_disconnected() -> bool:
            try:
                return await raw_request.is_disconnected()
            except Exception:
                return False

        # -- Session setup --
        session_id = body.session_id
        if session_id is None:
            sess = get_store().create_session()
            session_id = sess["id"]
        else:
            existing = get_store().get_session(session_id)
            if existing is None:
                yield f"event: error\ndata: {json.dumps({'error': 'Session not found'})}\n\n"
                return

        history_dicts = [
            {"role": m.role, "content": m.content} for m in body.history
        ]

        # Load session_state from DB
        saved_state_str = get_store().get_session_state(session_id)
        session_state_dict = json.loads(saved_state_str) if saved_state_str else None

        workflow = build_workflow()

        initial_state: dict = {
            "question": body.message,
            "workflow": [],
            "history": history_dicts,
            "_stream_answer": True,
            "source_mode": body.source_mode,
        }
        if history_dicts:
            initial_state["memory"] = {"history": list(history_dicts)}
        if session_state_dict:
            initial_state["session_state"] = SessionState.from_dict(session_state_dict)

        # -- Emit immediate start event (before any LLM call) --
        yield _sse_event("start", {"phase": "正在处理请求..."})

        # -- Stream workflow execution --
        final_state: dict | None = None
        answer_text: str | None = None  # captured from answer node for chunked delivery
        streamed_answer = ""
        did_stream_answer = False
        cancelled: bool = False
        try:
            async for event in workflow.astream(initial_state):
                # Check for client disconnect between nodes
                if await _is_disconnected():
                    cancelled = True
                    logger.info("Client disconnected during workflow execution")
                    break

                for node_name, state_update in event.items():
                    if node_name == "goal_analyzer":
                        goal = state_update.get("goal_analysis", {})
                        goal_type = goal.get("goal_type", "") if isinstance(goal, dict) else ""
                        yield _sse_event("thinking", {"phase": "分析用户目标", "goal_type": goal_type})
                    elif node_name == "planner":
                        yield _sse_event("planning", {"phase": "制定执行计划"})
                        yield _emit_task_update(state_update)
                    elif node_name == "knowledge":
                        yield _sse_event("searching", {"phase": "检索知识库"})
                    elif node_name == "query_rewriter":
                        yield _sse_event("searching", {"phase": "优化搜索查询"})
                        yield _emit_task_update(state_update)
                    elif node_name == "search":
                        yield _sse_event("searching", {"phase": "搜索网络信息"})
                        yield _emit_task_update(state_update)
                    elif node_name == "python":
                        yield _sse_event("executing", {"phase": "执行代码"})
                        yield _emit_task_update(state_update)
                    elif node_name == "tool_executor":
                        yield _emit_task_update(state_update)
                    elif node_name == "reflector":
                        yield _sse_event("thinking", {"phase": "检查执行结果"})
                        yield _emit_task_update(state_update)
                    elif node_name == "answer":
                        answer_text = state_update.get("answer", "")
                        yield _sse_event("generating", {"phase": "生成回答"})
                        yield _emit_task_update(state_update)
                        stream_messages = state_update.get("_answer_stream_messages")
                        if stream_messages and isinstance(stream_messages, list):
                            from agentflow.services.llm_service import get_llm_service

                            did_stream_answer = True
                            for token in get_llm_service().complete_stream(messages=stream_messages):
                                if await _is_disconnected():
                                    cancelled = True
                                    logger.info("Client disconnected during LLM streaming")
                                    break
                                if token:
                                    streamed_answer += token
                                    yield _sse_event("text", {"text": token})
                            answer_text = streamed_answer.strip()
                            if cancelled:
                                break
                    elif node_name == "memory":
                        final_state = dict(state_update)
                if cancelled:
                    break
        except Exception as exc:
            logger.exception("Streaming workflow failed")
            yield _sse_event("error", {"error": str(exc)})
            return

        # If cancelled, notify frontend and skip persistence
        if cancelled:
            yield _sse_event("cancelled", {"reason": "用户中断了对话"})
            return

        # -- Deliver answer text in chunks for true streaming feel --
        answer = answer_text or (final_state.get("answer", "") if final_state else "")
        if did_stream_answer:
            answer = answer.strip()
            if final_state is not None:
                final_state["answer"] = answer

        # Stream answer text incrementally
        if answer and not did_stream_answer:
            chunk_size = 15  # characters per chunk
            for i in range(0, len(answer), chunk_size):
                # Re-check disconnect during chunk delivery
                if await _is_disconnected():
                    logger.info("Client disconnected during chunk delivery")
                    yield _sse_event("cancelled", {"reason": "用户中断了对话"})
                    return
                chunk = answer[i:i + chunk_size]
                yield _sse_event("text", {"text": chunk})
                await asyncio.sleep(0.02)  # small delay for streaming effect

        # -- Persist messages --
        get_store().add_message("user", body.message, session_id=session_id)
        if answer:
            get_store().add_message("assistant", answer, session_id=session_id)

        # -- Persist session_state --
        if final_state:
            ctx = WorkflowContext(final_state)
            result_dict = ctx.to_dict()
            new_state = result_dict.get("session_state")
            if new_state and isinstance(new_state, dict):
                get_store().update_session_state(session_id, json.dumps(new_state, ensure_ascii=False))

        # -- Auto-title --
        sess = get_store().get_session(session_id)
        if sess and sess["title"] == "新对话":
            title = body.message[:50]
            if len(body.message) > 50:
                title += "…"
            get_store().update_session_title(session_id, title)

        done_data: dict[str, object] = {"answer": answer, "session_id": session_id}
        if final_state and final_state.get("_degraded"):
            done_data["degraded"] = True
            done_data["degraded_reason"] = str(final_state.get("_llm_error", "unknown"))
        # Final task queue sync — ensure frontend shows correct final state
        if final_state:
            final_tasks = _emit_task_update(final_state)
            if final_tasks:
                yield final_tasks
        yield _sse_event("done", done_data)

    return StreamingResponse(_event_generator(), media_type="text/event-stream")


def _sse_event(event: str, data: dict) -> str:
    """Format an SSE event string."""
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


def _emit_task_update(state_update: dict) -> str:
    """Extract task queue from workflow state and return an SSE ``task_update`` event.

    Returns empty string when the task_queue is empty.
    """
    queue = state_update.get("task_queue", []) or []
    if not queue:
        return ""
    tasks = [
        {
            "id": t.get("task_id", ""),
            "title": t.get("title", ""),
            "tool": t.get("tool", ""),
            "status": t.get("status", "todo"),
        }
        for t in queue
    ]
    return _sse_event("task_update", {"tasks": tasks})


# -- Knowledge base: document management ------------------------------------


@router.post("/upload")
async def upload_file(file: UploadFile = File(...)) -> JSONResponse:
    """Upload a document, parse it, and index it into the knowledge base."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    # Validate file type
    allowed_types = {
        ".pdf", ".docx", ".doc", ".txt", ".md", ".markdown",
        ".html", ".htm", ".xlsx", ".xls", ".pptx", ".csv", ".epub",
        ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".go", ".rs",
        ".c", ".cpp", ".h", ".hpp", ".zip",
    }
    ext = Path(file.filename).suffix.lower()
    if ext not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. "
                   f"Allowed: PDF, DOCX, TXT, MD, HTML, XLSX, PPTX, CSV, EPUB, 代码文件",
        )

    # Handle ZIP archives: extract and ingest each file
    if ext == ".zip":
        return await _handle_zip_upload(file)

    # Save uploaded file temporarily
    safe_filename = Path(file.filename).name
    temp_path = UPLOAD_DIR / safe_filename
    try:
        content = await file.read()
        temp_path.write_bytes(content)
        logger.info("Saved uploaded file: %s (%d bytes)", safe_filename, len(content))

        # Ingest into knowledge base
        doc_id = get_knowledge_store().add_document(temp_path, safe_filename)
        return JSONResponse(
            content={
                "status": "ok",
                "document_id": doc_id,
                "filename": safe_filename,
                "size": len(content),
            }
        )
    except Exception as exc:
        logger.exception("Upload ingestion failed for %s", safe_filename)
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        # Clean up temp file after indexing
        if temp_path.exists():
            temp_path.unlink(missing_ok=True)


async def _handle_zip_upload(file: UploadFile) -> JSONResponse:
    """Extract a ZIP archive and ingest each contained document."""
    import zipfile
    from agentflow.knowledge.parser import _read_raw_from_bytes

    content = await file.read()
    total = 0
    success = 0
    failed: list[str] = []
    allowed_exts = {
        ".pdf", ".docx", ".doc", ".txt", ".md", ".markdown",
        ".html", ".htm", ".xlsx", ".xls", ".pptx", ".csv", ".epub",
        ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".go", ".rs",
        ".c", ".cpp", ".h", ".hpp",
    }

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        for info in zf.infolist():
            fname = Path(info.filename)
            if info.filename.startswith("__MACOSX/") or info.filename.startswith("."):
                continue
            if fname.suffix.lower() not in allowed_exts:
                continue
            total += 1
            try:
                raw = zf.read(info.filename)
                file_type = fname.suffix.lstrip(".").lower()
                text = _read_raw_from_bytes(raw, file_type)
                if not text.strip():
                    continue
                # Save to temp file and ingest
                temp_path = UPLOAD_DIR / fname.name
                temp_path.write_bytes(raw)
                try:
                    get_knowledge_store().add_document(temp_path, fname.name)
                    success += 1
                finally:
                    if temp_path.exists():
                        temp_path.unlink(missing_ok=True)
            except Exception as exc:
                logger.warning("ZIP entry failed: %s (%s)", info.filename, exc)
                failed.append(info.filename)

    return JSONResponse(content={
        "status": "ok",
        "total": total,
        "success": success,
        "failed": failed,
        "filename": file.filename,
        "size": len(content),
    })


@router.get("/knowledge/documents")
def list_documents() -> list[dict[str, object]]:
    """List all indexed documents."""
    return get_knowledge_store().list_documents()


@router.delete("/knowledge/documents/{doc_id}")
def delete_document(doc_id: int) -> JSONResponse:
    """Delete a document and its chunks/embeddings from the knowledge base."""
    get_knowledge_store().delete_document(doc_id)
    return JSONResponse(content={"status": "deleted", "document_id": doc_id})


@router.post("/knowledge/search")
def search_knowledge(query: str, top_k: int = 5) -> list[dict[str, object]]:
    """Search the knowledge base for relevant chunks."""
    return get_knowledge_store().search(query, top_k=top_k)


# -- Chat history ----------------------------------------------------------


@router.get("/history")
def history(limit: int = 20) -> list[dict[str, str]]:
    """Fetch recent chat history."""
    return get_store().list_messages(limit=limit)


# -- File operations (agent-generated files) --------------------------------


class CreateFileRequest(BaseModel):
    filename: str
    content: str
    workspace_path: str | None = None


class ReadFileRequest(BaseModel):
    path: str
    workspace_path: str | None = None


def _safe_relative_file_path(raw_path: str) -> Path:
    """Return a safe relative file path, preserving subdirectories."""
    rel = Path(raw_path.replace("\\", "/"))
    if rel.is_absolute() or any(part in {"", ".", ".."} for part in rel.parts):
        raise HTTPException(status_code=400, detail="Invalid file path")
    return rel


def _resolve_generated_file(path: str, workspace_path: str | None = None) -> Path:
    """Resolve a generated file path under outputs/ or the active workspace."""
    base = _resolve_workspace_child(workspace_path) if workspace_path else OUTPUT_DIR.resolve()
    raw = Path(path)
    target = raw.resolve() if raw.is_absolute() else (base / _safe_relative_file_path(path)).resolve()
    if not _is_relative_to(target, base):
        raise HTTPException(status_code=400, detail="File path is outside the workspace")
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail=f"File not found: {path}")
    return target


@router.post("/files/create")
def create_file(req: CreateFileRequest) -> JSONResponse:
    """Write a proposed file to the outputs/ or workspace directory."""
    rel_path = _safe_relative_file_path(req.filename)

    if req.workspace_path:
        base = _resolve_workspace_child(req.workspace_path)
        if not base.exists() or not base.is_dir():
            raise HTTPException(status_code=400, detail="Invalid workspace path")
        target = (base / rel_path).resolve()
    else:
        base = OUTPUT_DIR.resolve()
        target = (base / rel_path).resolve()

    if not _is_relative_to(target, base):
        raise HTTPException(status_code=400, detail="File path is outside the workspace")

    if target.exists():
        raise HTTPException(status_code=409, detail="File already exists")
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(req.content, encoding="utf-8")
    return JSONResponse(
        content={
            "status": "created",
            "filename": rel_path.as_posix(),
            "path": str(target),
        }
    )


@router.post("/files/read")
def read_generated_file(req: ReadFileRequest) -> JSONResponse:
    """Read a generated file for in-app preview."""
    target = _resolve_generated_file(req.path, req.workspace_path)
    max_bytes = 1024 * 1024
    data = target.read_bytes()
    truncated = len(data) > max_bytes
    if target.suffix.lower() == ".docx":
        text = _read_docx_preview(target)
        truncated = False
    else:
        text = data[:max_bytes].decode("utf-8", errors="replace")
    return JSONResponse(
        content={
            "filename": target.name,
            "path": str(target),
            "content": text,
            "truncated": truncated,
            "size": len(data),
        }
    )


def _read_docx_preview(path: Path) -> str:
    """Extract readable text from a .docx file for preview."""
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="python-docx is not installed") from exc

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to read docx file: {exc}") from exc

    parts: list[str] = []
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if paragraphs:
        parts.extend(paragraphs)

    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"\n[表格 {table_index}]")
            parts.extend(rows)

    return "\n\n".join(parts) if parts else "（该 Word 文档没有可预览的文本内容）"


@router.get("/files")
def list_output_files(workspace_path: str | None = None) -> list[dict[str, object]]:
    """List files in the outputs/ directory or a workspace path."""
    if workspace_path:
        base = _resolve_workspace_child(workspace_path)
        if not base.exists() or not base.is_dir():
            raise HTTPException(status_code=400, detail="Invalid workspace path")
    else:
        base = OUTPUT_DIR

    if not base.exists():
        return []
    files: list[dict[str, object]] = []
    for p in sorted(base.rglob("*")):
        if p.is_file():
            rel = p.relative_to(base)
            files.append(
                {
                    "filename": str(rel),
                    "size": p.stat().st_size,
                    "created_at": datetime.fromtimestamp(p.stat().st_ctime).isoformat(),
                    "path": str(p),
                }
            )
    return files


# -- Workspace operations ----------------------------------------------------


@router.get("/workspace")
def get_workspace() -> JSONResponse:
    """Check if a path is a valid workspace directory."""
    return JSONResponse(content={"status": "ok", "message": "Provide a path via POST /workspace/set"})


class SetWorkspaceRequest(BaseModel):
    path: str


@router.post("/workspace/set")
def set_workspace(req: SetWorkspaceRequest) -> JSONResponse:
    """Validate and set workspace folder path."""
    global _workspace_root
    p = Path(req.path).resolve()
    if not p.exists():
        raise HTTPException(status_code=404, detail=f"Path does not exist: {req.path}")
    if not p.is_dir():
        raise HTTPException(status_code=400, detail="Path is not a directory")
    # Test write permission
    test_file = p / ".omni_forge_write_test"
    try:
        test_file.write_text("test", encoding="utf-8")
        test_file.unlink()
    except OSError as exc:
        raise HTTPException(status_code=403, detail=f"No write permission: {exc}")
    _workspace_root = p
    return JSONResponse(content={"status": "ok", "path": str(p)})


class CreateFolderRequest(BaseModel):
    parent_path: str
    folder_name: str


@router.post("/workspace/create-folder")
def create_workspace_folder(req: CreateFolderRequest) -> JSONResponse:
    """Create a new folder under the given parent path."""
    parent = _resolve_workspace_child(req.parent_path)
    if not parent.exists() or not parent.is_dir():
        raise HTTPException(status_code=404, detail=f"Parent path does not exist: {req.parent_path}")
    safe_name = Path(req.folder_name).name
    if not safe_name or safe_name in {".", ".."}:
        raise HTTPException(status_code=400, detail="Invalid folder name")
    target = parent / safe_name
    if target.exists():
        raise HTTPException(status_code=409, detail=f"Folder already exists: {safe_name}")
    try:
        target.mkdir(parents=True, exist_ok=False)
    except OSError as exc:
        raise HTTPException(status_code=500, detail=f"Failed to create folder: {exc}")
    return JSONResponse(content={"status": "created", "path": str(target)})


@router.get("/workspace/browse")
def browse_directory(path: str = ".") -> JSONResponse:
    """List directories and files at the given path for folder browsing."""
    base = _resolve_workspace_child(path)
    if not base.exists() or not base.is_dir():
        raise HTTPException(status_code=404, detail=f"Path does not exist: {path}")
    entries: list[dict[str, object]] = []
    for p in sorted(base.iterdir()):
        entries.append({
            "name": p.name,
            "is_dir": p.is_dir(),
            "path": str(p),
        })
    return JSONResponse(content={"current_path": str(base), "entries": entries})


# -- Sessions ----------------------------------------------------------------


@router.post("/sessions/create")
def create_session() -> JSONResponse:
    """Create a new chat session."""
    sess = get_store().create_session()
    return JSONResponse(content=sess)


@router.get("/sessions")
def list_sessions(limit: int = 50) -> list[dict[str, object]]:
    """List all chat sessions, most recent first."""
    return get_store().list_sessions(limit=limit)


@router.get("/sessions/{session_id}/messages")
def get_session_messages(session_id: int) -> list[dict[str, object]]:
    """Get all messages for a session."""
    sess = get_store().get_session(session_id)
    if sess is None:
        raise HTTPException(status_code=404, detail="Session not found")
    return get_store().get_session_messages(session_id)


@router.put("/sessions/{session_id}/rename")
def rename_session(session_id: int, body: dict[str, str]) -> JSONResponse:
    """Rename a session."""
    title = body.get("title", "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title is required")
    ok = get_store().update_session_title(session_id, title)
    if not ok:
        raise HTTPException(status_code=404, detail="Session not found")
    return JSONResponse(content={"status": "ok"})


@router.delete("/sessions/{session_id}")
def delete_session_endpoint(session_id: int) -> JSONResponse:
    """Delete a session and all its messages."""
    ok = get_store().delete_session(session_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Session not found")
    return JSONResponse(content={"status": "deleted"})


@router.post("/sessions/cleanup")
def cleanup_sessions() -> JSONResponse:
    """Manually trigger cleanup of expired sessions and memories.

    Uses settings ``session_ttl_hours`` and ``memory_ttl_days``.
    """
    from agentflow.config.settings import settings
    deleted_sessions = get_store().delete_sessions_older_than(settings.session_ttl_hours)
    deleted_memories = get_store().delete_old_memories(settings.memory_ttl_days)
    logger.info(
        "Manual cleanup: removed %d sessions, %d memory entries",
        deleted_sessions, deleted_memories,
    )
    return JSONResponse(content={
        "status": "ok",
        "deleted_sessions": deleted_sessions,
        "deleted_memories": deleted_memories,
    })


# -- Model configuration ---------------------------------------------------


@router.get("/models")
def list_models() -> list[dict[str, object]]:
    """List all configured LLM models (API key excluded in response)."""
    from agentflow.models.model_config import LLMModelResponse
    rows = get_store().get_all_models()
    return [LLMModelResponse.from_db_row(r).model_dump() for r in rows]


@router.post("/models", status_code=201)
def create_model(config: LLMModelCreate) -> dict[str, object]:
    """Create a new LLM model configuration."""
    model_id = get_store().add_model(
        name=config.name,
        provider=config.provider,
        base_url=config.base_url,
        api_key=config.api_key,
        model_name=config.model_name,
        temperature=config.temperature,
        max_tokens=config.max_tokens,
    )
    return {"id": model_id, "status": "created"}


@router.put("/models/{model_id}")
def update_model(model_id: int, config: LLMModelUpdate) -> dict[str, object]:
    """Update an existing LLM model configuration."""
    ok = get_store().update_model(model_id, **config.model_dump(exclude_unset=True))
    if not ok:
        raise HTTPException(status_code=404, detail="Model not found")
    return {"status": "updated"}


@router.delete("/models/{model_id}")
def delete_model(model_id: int) -> dict[str, object]:
    """Delete an LLM model configuration."""
    ok = get_store().delete_model(model_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Model not found")
    return {"status": "deleted"}


@router.post("/models/{model_id}/activate")
def activate_model(model_id: int) -> dict[str, object]:
    """Set a model as the active LLM configuration."""
    model = get_store().get_model(model_id)
    if model is None:
        raise HTTPException(status_code=404, detail="Model not found")
    get_store().set_active_model(model_id)
    return {"status": "activated", "model_name": model["model_name"]}


# -- Long-term memory ------------------------------------------------------


@router.get("/memory")
def list_memories(category: str = "", limit: int = 50) -> list[dict[str, object]]:
    """List all long-term memories, optionally filtered by category."""
    from agentflow.services.long_term_memory import LongTermMemory
    return LongTermMemory(db=get_store()).get_all(category=category)


@router.get("/memory/search")
def search_memories(query: str, limit: int = 10) -> list[dict[str, object]]:
    """Search long-term memories by keyword."""
    from agentflow.services.long_term_memory import LongTermMemory
    return LongTermMemory(db=get_store()).recall(query, limit=limit)


@router.delete("/memory/{key}")
def delete_memory(key: str) -> JSONResponse:
    """Delete a specific long-term memory."""
    from agentflow.services.long_term_memory import LongTermMemory
    ok = LongTermMemory(db=get_store()).forget(key)
    if not ok:
        raise HTTPException(status_code=404, detail="Memory not found")
    return JSONResponse(content={"status": "deleted"})


@router.delete("/memory")
def clear_memories(category: str = "") -> JSONResponse:
    """Clear all long-term memories, optionally filtered by category."""
    from agentflow.services.long_term_memory import LongTermMemory
    LongTermMemory(db=get_store()).clear(category=category)
    return JSONResponse(content={"status": "cleared"})


# -- Tool introspection -------------------------------------------------------


@router.get("/tools")
def list_tools() -> list[dict[str, object]]:
    """List all registered tools with metadata."""
    from agentflow.graph.workflow import get_executor
    ex = get_executor()
    if ex is None:
        return []
    return ex.tool_metadata()


@router.get("/tools/capabilities")
def list_tool_capabilities() -> list[str]:
    """List all aggregated capabilities from registered tools."""
    from agentflow.graph.workflow import get_executor
    ex = get_executor()
    if ex is None:
        return []
    return ex.get_capabilities()


@router.get("/tools/executor")
def executor_status() -> dict[str, object]:
    """Return the Executor's status summary."""
    from agentflow.graph.workflow import get_executor
    ex = get_executor()
    if ex is None:
        return {"status": "not_initialised"}
    return {
        "status": "ready",
        "tools": ex.list_tools(),
        "capabilities": ex.get_capabilities(),
        "summary": ex.summary,
    }
